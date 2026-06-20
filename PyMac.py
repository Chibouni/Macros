import pandas as pd
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Border, Side, Alignment


# ==========================================================
# EXCEL DASHBOARD CREATION
# ==========================================================

def create_dashboard(data, output_excel):

    df = pd.DataFrame(data)

    # Normalize status
    df["Status"] = df["Status"].fillna("").str.upper().str.strip()

        # Failed tests → keep issue
    nok_df = df[df["Status"] == "NOK"][
        ["Test_ID", "Status", "Issue"]
    ]

    # Passed tests → discard comment + issue
    ok_df = df[df["Status"] == "OK"][
        ["Test_ID", "Status"]
    ]

    # Not Processed → keep BOTH issue + comment
    np_df = df[df["Status"] == "NP"][
        ["Test_ID", "Status", "Comment", "Issue", ]
    ]

    # Missing status
    empty_df = df[df["Status"] == ""][
        ["Test_ID"]
    ]


    wb = Workbook()
    ws = wb.active
    ws.title = "Test Dashboard"


    # ==========================
    # STYLES
    # ==========================

    white_font = Font(color="FFFFFF", bold=True)
    title_font = Font(size=14, bold=True)

    border = Border(
        left=Side(style="thin"),
        right=Side(style="thin"),
        top=Side(style="thin"),
        bottom=Side(style="thin")
    )


    def apply_header_style(cell, color):
        cell.fill = PatternFill(
            start_color=color,
            end_color=color,
            fill_type="solid"
        )
        cell.font = white_font
        cell.border = border
        cell.alignment = Alignment(horizontal="center")


    # ==========================
    # SUMMARY SECTION
    # ==========================

    ws["A2"] = "TEST EXECUTION SUMMARY"
    ws["A2"].font = title_font

    ws["A4"] = "Metric"
    ws["B4"] = "Count"

    apply_header_style(ws["A4"], "404040")
    apply_header_style(ws["B4"], "404040")


    summary = [
        ("Passed Tests (OK)", len(ok_df)),
        ("Failed Tests (NOK)", len(nok_df)),
        ("Not Processed (NP)", len(np_df)),
        ("Missing Status", len(empty_df)),
        ("Total Tests", len(df))
    ]

    row = 5
    for metric, value in summary:

        ws[f"A{row}"] = metric
        ws[f"B{row}"] = value

        ws[f"A{row}"].border = border
        ws[f"B{row}"].border = border

        row += 1


    # ==========================
    # TABLE WRITER
    # ==========================

    def write_table(start_row, start_col, dataframe, title, color):

        # Title
        ws.cell(start_row, start_col, title)
        ws.cell(start_row, start_col).font = Font(bold=True)

        # Headers
        for c, col_name in enumerate(dataframe.columns):

            cell = ws.cell(
                row=start_row + 1,
                column=start_col + c,
                value=col_name
            )

            apply_header_style(cell, color)

        # Data rows
        for r, row_data in enumerate(dataframe.values):

            for c, value in enumerate(row_data):

                cell = ws.cell(
                    row=start_row + 2 + r,
                    column=start_col + c,
                    value=value
                )

                cell.border = border


    # ==========================
    # WRITE TABLES
    # ==========================

    write_table(
        start_row=11,
        start_col=1,
        dataframe=nok_df,
        title="Failed Tests (NOK)",
        color="8C2D19"      # dark red
    )


    write_table(
        start_row=11,
        start_col=5,
        dataframe=ok_df,
        title="Passed Tests (OK)",
        color="2E6930"      # green
    )


    write_table(
        start_row=11,
        start_col=8,
        dataframe=np_df,
        title="Not Processed (NP)",
        color="4A6572"      # blue
    )


    write_table(
        start_row=11,
        start_col=13,
        dataframe=empty_df,
        title="Missing Status",
        color="808080"      # gray
    )


    # ==========================
    # AUTO COLUMN WIDTH
    # ==========================

    for column in ws.columns:

        max_length = 0
        column_letter = column[0].column_letter

        for cell in column:

            try:
                if cell.value:
                    max_length = max(
                        max_length,
                        len(str(cell.value))
                    )
            except:
                pass

        ws.column_dimensions[column_letter].width = max_length + 3


    wb.save(output_excel)

    print("Dashboard created:", output_excel)


# ==========================================================
# WORD EXTRACTION
# ==========================================================

def extract_test_data(docx_path, output_excel):

    doc = Document(docx_path)

    data = []

    # Tables come in pairs
    for i in range(0, len(doc.tables), 2):

        if i + 1 >= len(doc.tables):
            break

        header_table = doc.tables[i]
        details_table = doc.tables[i + 1]

        # ----------------------------------------
        # Extract Test ID
        # ----------------------------------------

        raw_header = header_table.cell(0, 0).text

        # Example:
        # SWTR_D0000958_00001 - something
        test_id = raw_header.split(" - ")[0].strip()


        # ----------------------------------------
        # Find Comments/Status row
        # ----------------------------------------

        data_row_index = -1

        for row_idx, row in enumerate(details_table.rows):

            row_text = [
                cell.text.strip().lower()
                for cell in row.cells
            ]

            if "comments" in row_text and "status" in row_text:

                data_row_index = row_idx + 1
                break


        # ----------------------------------------
        # Extract data
        # ----------------------------------------

        if data_row_index != -1 and data_row_index < len(details_table.rows):

            row = details_table.rows[data_row_index]

            try:

                comment = row.cells[0].text.strip()

                status = row.cells[1].text.strip().upper()

                issue = ""

                if len(row.cells) > 2:
                    issue = row.cells[2].text.strip()


                data.append({
                    "Test_ID": test_id,
                    "Status": status,
                    "Comment": comment,
                    "Issue": issue
                })

            except IndexError:
                continue


    print("Extracted tests:", len(data))


    # Create styled dashboard
    create_dashboard(data, output_excel)


# ==========================================================
# MAIN
# ==========================================================

if __name__ == "__main__":

    extract_test_data(
        "D0000958 SWTP A1 _ OVERALL validation test report for SPU70-66-GEN1.5 CS.docx",
        "Test_Dashboard.xlsx"
    )