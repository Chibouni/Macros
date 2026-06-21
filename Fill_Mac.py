from docx import Document


def fill_empty_fields(input_docx,
                      output_docx,
                      default_comment="Not executed",
                      default_status="NP"):

    doc = Document(input_docx)

    updated_count = 0

    # tables are in pairs
    for i in range(0, len(doc.tables), 2):

        if i + 1 >= len(doc.tables):
            break

        details_table = doc.tables[i + 1]

        data_row_index = -1

        # ---------------------------------------
        # find header row
        # ---------------------------------------
        for row_idx, row in enumerate(details_table.rows):

            row_text = [
                cell.text.strip().lower()
                for cell in row.cells
            ]

            if "comments" in row_text and "status" in row_text:

                data_row_index = row_idx + 1
                break

        # ---------------------------------------
        # modify data row
        # ---------------------------------------
        if data_row_index != -1 and data_row_index < len(details_table.rows):

            row = details_table.rows[data_row_index]

            comment_cell = row.cells[0]
            status_cell = row.cells[1]

            comment = comment_cell.text.strip()
            status = status_cell.text.strip()

            # fill empty comment
            if comment == "":
                comment_cell.text = default_comment
                updated_count += 1

            # fill empty status
            if status == "":
                status_cell.text = default_status
                updated_count += 1


    doc.save(output_docx)

    print("Updated fields:", updated_count)
    print("Saved:", output_docx)


if __name__ == "__main__":

    fill_empty_fields(
        input_docx="EX1.docx",
        output_docx="updated_report.docx",
        default_comment="Not executed",
        default_status="NP"
    )